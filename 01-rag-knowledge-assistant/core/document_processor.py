"""
Advanced document processing engine for RAG systems
Secure file handling with magic number validation and sandboxed processing
"""

import os
import hashlib
import mimetypes
import tempfile
import shutil
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import asyncio
import aiofiles
import logging

import PyPDF2
import docx
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import tiktoken

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent))
from simple_config import get_rag_config

config = get_rag_config()
logger = logging.getLogger(__name__)

# Security constants
ALLOWED_MIME_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'txt': 'text/plain',
    'md': 'text/markdown',
    'csv': 'text/csv',
    'json': 'application/json',
    'py': 'text/x-python',
    'js': 'application/javascript',
    'ts': 'application/typescript',
    'html': 'text/html'
}

MAGIC_NUMBERS = {
    'pdf': [b'%PDF-'],
    'docx': [b'PK\x03\x04'],
    'txt': [],  # No magic number for plain text
    'md': [],   # No magic number for markdown
    'csv': [],  # No magic number for CSV
    'json': [],  # No magic number for JSON
    'py': [],   # No magic number for Python
    'js': [],   # No magic number for JavaScript
    'ts': [],   # No magic number for TypeScript
    'html': [b'<!DOCTYPE html', b'<html', b'<HTML']
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


class ProcessingError(Exception):
    """Exception raised during document processing"""
    pass


@dataclass
class DocumentChunk:
    """Represents a processed document chunk with metadata"""
    id: str
    content: str
    metadata: Dict[str, Any]
    chunk_index: int = 0
    total_chunks: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "content": self.content,
            "metadata": self.metadata,
            "chunk_index": self.chunk_index,
            "total_chunks": self.total_chunks
        }


@dataclass
class ProcessedDocument:
    """Container for processed document with all chunks"""
    id: str
    source_path: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processing_stats: Dict[str, Any]
    
    @property
    def total_chunks(self) -> int:
        return len(self.chunks)
    
    @property
    def total_tokens(self) -> int:
        return sum(chunk.metadata.get("token_count", 0) for chunk in self.chunks)


class SecureFileHandler:
    """Secure file handling with validation and sandboxing"""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.mkdtemp(prefix="secure_docs_"))
        self.temp_dir.chmod(0o700)  # Owner read/write/execute only
        
    def __del__(self):
        """Clean up temporary directory"""
        if hasattr(self, 'temp_dir') and self.temp_dir.exists():
            shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def validate_file_security(self, file_path: str) -> bool:
        """Comprehensive file security validation"""
        file_path = Path(file_path)
        
        # Check file exists
        if not file_path.exists():
            raise ValueError(f"File not found: {file_path}")
        
        # Check file size
        if file_path.stat().st_size > MAX_FILE_SIZE:
            raise ValueError(f"File too large: {file_path.stat().st_size} bytes > {MAX_FILE_SIZE}")
        
        # Check file extension
        extension = file_path.suffix.lower().lstrip('.')
        if extension not in ALLOWED_MIME_TYPES:
            raise ValueError(f"File type not allowed: {extension}")
        
        # Use mimetypes for basic validation (less secure but works without libmagic)
        detected_mime, _ = mimetypes.guess_type(str(file_path))
        expected_mime = ALLOWED_MIME_TYPES.get(extension)
        
        if detected_mime != expected_mime and expected_mime:
            # Allow text files to have various MIME types
            if extension in ['txt', 'md', 'py', 'js', 'ts', 'csv', 'json']:
                if detected_mime and not detected_mime.startswith('text/'):
                    logger.warning(f"MIME type mismatch: {detected_mime} vs {expected_mime}")
            else:
                logger.warning(f"MIME type mismatch: {detected_mime} vs {expected_mime}")
        
        # Validate magic numbers
        if extension in MAGIC_NUMBERS and MAGIC_NUMBERS[extension]:
            with open(file_path, 'rb') as f:
                file_header = f.read(1024)
            
            valid_magic = any(
                file_header.startswith(magic_num) 
                for magic_num in MAGIC_NUMBERS[extension]
            )
            
            if not valid_magic:
                raise ValueError(f"Invalid file format: {extension}")
        
        return True
    
    def create_secure_copy(self, source_path: str) -> str:
        """Create secure copy of file in sandboxed directory"""
        source_path = Path(source_path)
        
        # Generate secure filename
        secure_name = hashlib.sha256(
            f"{source_path.name}{datetime.utcnow().isoformat()}".encode()
        ).hexdigest()[:16]
        
        secure_path = self.temp_dir / f"{secure_name}{source_path.suffix}"
        
        # Copy file with restricted permissions
        shutil.copy2(source_path, secure_path)
        secure_path.chmod(0o600)  # Owner read/write only
        
        return str(secure_path)
    
    def sanitize_content(self, content: str) -> str:
        """Sanitize file content"""
        # Remove potentially dangerous characters
        dangerous_chars = ['\x00', '\x01', '\x02', '\x03', '\x04', '\x05', '\x06', '\x07', '\x08']
        for char in dangerous_chars:
            content = content.replace(char, '')
        
        # Limit content length
        if len(content) > 10_000_000:  # 10MB text limit
            content = content[:10_000_000]
            logger.warning("Content truncated due to size limit")
        
        return content


class DocumentProcessor:
    """Advanced document processing with security-first approach"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.file_handler = SecureFileHandler()
        
        # Initialize format handlers
        self.format_handlers = {
            "pdf": self._process_pdf,
            "docx": self._process_docx,
            "txt": self._process_text,
            "md": self._process_markdown,
            "csv": self._process_csv,
            "json": self._process_json,
            "py": self._process_code,
            "js": self._process_code,
            "ts": self._process_code,
            "html": self._process_html
        }
        
        logger.info("Document processor initialized with security validation")
    
    async def process_document(self, file_path: str, 
                              metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """Process a document with comprehensive security validation"""
        try:
            # Validate file security
            self.file_handler.validate_file_security(file_path)
            
            # Create secure copy
            secure_path = self.file_handler.create_secure_copy(file_path)
            
            # Extract base metadata
            base_metadata = await self._extract_base_metadata(secure_path)
            if metadata:
                base_metadata.update(metadata)
            
            # Process document content
            file_format = self._detect_format(secure_path)
            if file_format not in self.format_handlers:
                raise ValueError(f"Unsupported file format: {file_format}")
            
            handler = self.format_handlers[file_format]
            raw_content, extracted_metadata = await handler(secure_path)
            
            # Sanitize content
            raw_content = self.file_handler.sanitize_content(raw_content)
            
            # Merge metadata
            final_metadata = {**base_metadata, **extracted_metadata}
            
            # Create document chunks
            chunks = await self._create_chunks(
                content=raw_content,
                source_path=file_path,  # Use original path for metadata
                base_metadata=final_metadata
            )
            
            # Calculate processing stats
            processing_stats = {
                "total_chunks": len(chunks),
                "total_characters": len(raw_content),
                "total_tokens": sum(chunk.metadata.get("token_count", 0) for chunk in chunks),
                "processing_time": datetime.utcnow().isoformat(),
                "file_format": file_format,
                "file_size_bytes": Path(file_path).stat().st_size,
                "security_validated": True
            }
            
            # Create document ID
            doc_id = self._generate_document_id(file_path, raw_content)
            
            return ProcessedDocument(
                id=doc_id,
                source_path=file_path,
                chunks=chunks,
                metadata=final_metadata,
                processing_stats=processing_stats
            )
            
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            raise
    
    async def _extract_base_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic file metadata"""
        file_path = Path(file_path)
        stat = file_path.stat()
        
        return {
            "source_file": file_path.name,
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "mime_type": magic.from_file(str(file_path), mime=True),
            "processed_at": datetime.utcnow().isoformat(),
            "processor_version": "2.0.0"
        }
    
    def _detect_format(self, file_path: str) -> str:
        """Detect file format from extension and validation"""
        extension = Path(file_path).suffix.lower().lstrip('.')
        
        if extension in ALLOWED_MIME_TYPES:
            return extension
        
        raise ValueError(f"Unsupported file format: {extension}")
    
    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF with security validation"""
        try:
            content = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Check for password protection
                if reader.is_encrypted:
                    raise ValueError("Password-protected PDFs not supported")
                
                # Limit page count for security
                max_pages = min(len(reader.pages), 1000)
                
                for page_num in range(max_pages):
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text:
                        content += page_text + "\n\n"
                
                metadata = {
                    "page_count": len(reader.pages),
                    "pages_processed": max_pages,
                    "document_info": reader.metadata or {}
                }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    async def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX with security validation"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "document_type": "docx"
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    async def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process plain text with encoding detection"""
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            
            for encoding in encodings:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        content = await f.read()
                    
                    metadata = {
                        "encoding": encoding,
                        "line_count": len(content.split('\n')),
                        "document_type": "text"
                    }
                    
                    return content, metadata
                    
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode text file")
            
        except Exception as e:
            logger.error(f"Text processing failed: {str(e)}")
            raise ValueError(f"Failed to process text: {str(e)}")
    
    async def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown files"""
        content, metadata = await self._process_text(file_path)
        
        # Count markdown elements
        lines = content.split('\n')
        headers = len([line for line in lines if line.strip().startswith('#')])
        links = content.count('[') + content.count('](')
        
        metadata.update({
            "document_type": "markdown",
            "header_count": headers,
            "link_count": links
        })
        
        return content, metadata
    
    async def _process_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process CSV files with security limits"""
        try:
            # Limit file size for CSV processing
            if Path(file_path).stat().st_size > 10_000_000:  # 10MB limit
                raise ValueError("CSV file too large")
            
            df = pd.read_csv(file_path, nrows=10000)  # Limit rows
            
            # Convert to text representation
            content = df.to_string(index=False)
            
            metadata = {
                "document_type": "csv",
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist()
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"CSV processing failed: {str(e)}")
            raise ValueError(f"Failed to process CSV: {str(e)}")
    
    async def _process_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process JSON files with validation"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Validate JSON structure
            import json
            try:
                data = json.loads(content)
                
                # Convert to readable format
                content = json.dumps(data, indent=2, ensure_ascii=False)
                
                metadata = {
                    "document_type": "json",
                    "json_valid": True,
                    "data_type": type(data).__name__
                }
                
                if isinstance(data, dict):
                    metadata["key_count"] = len(data)
                elif isinstance(data, list):
                    metadata["item_count"] = len(data)
                
            except json.JSONDecodeError:
                # If invalid JSON, treat as text
                metadata = {
                    "document_type": "json",
                    "json_valid": False
                }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"JSON processing failed: {str(e)}")
            raise ValueError(f"Failed to process JSON: {str(e)}")
    
    async def _process_code(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process code files with syntax awareness"""
        content, metadata = await self._process_text(file_path)
        
        # Basic code analysis
        lines = content.split('\n')
        code_lines = len([line for line in lines if line.strip() and not line.strip().startswith('#')])
        comment_lines = len([line for line in lines if line.strip().startswith('#')])
        
        # Detect language from extension
        extension = Path(file_path).suffix.lower()
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript'
        }
        
        metadata.update({
            "document_type": "code",
            "language": language_map.get(extension, "unknown"),
            "code_lines": code_lines,
            "comment_lines": comment_lines,
            "total_lines": len(lines)
        })
        
        return content, metadata
    
    async def _process_html(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process HTML files with tag stripping"""
        try:
            from bs4 import BeautifulSoup
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                html_content = await f.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in content.splitlines())
            content = '\n'.join(line for line in lines if line)
            
            metadata = {
                "document_type": "html",
                "title": soup.title.string if soup.title else None,
                "has_links": len(soup.find_all('a')) > 0,
                "has_images": len(soup.find_all('img')) > 0
            }
            
            return content, metadata
            
        except ImportError:
            # Fallback without BeautifulSoup
            content, metadata = await self._process_text(file_path)
            metadata["document_type"] = "html"
            return content, metadata
        
        except Exception as e:
            logger.error(f"HTML processing failed: {str(e)}")
            raise ValueError(f"Failed to process HTML: {str(e)}")
    
    async def _create_chunks(self, content: str, source_path: str, 
                           base_metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Create document chunks with token counting"""
        if not content.strip():
            return []
        
        # Split content into chunks
        chunks = self.text_splitter.split_text(content)
        
        document_chunks = []
        for i, chunk_text in enumerate(chunks):
            # Count tokens
            token_count = len(self.tokenizer.encode(chunk_text))
            
            # Create chunk metadata
            chunk_metadata = {
                **base_metadata,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "token_count": token_count,
                "character_count": len(chunk_text),
                "chunk_type": "text"
            }
            
            # Generate chunk ID
            chunk_id = hashlib.md5(
                f"{source_path}_{i}_{chunk_text[:100]}".encode()
            ).hexdigest()
            
            document_chunks.append(DocumentChunk(
                id=chunk_id,
                content=chunk_text,
                metadata=chunk_metadata,
                chunk_index=i,
                total_chunks=len(chunks)
            ))
        
        return document_chunks
    
    def _generate_document_id(self, file_path: str, content: str) -> str:
        """Generate unique document ID"""
        content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
        file_hash = hashlib.md5(file_path.encode()).hexdigest()[:8]
        return f"{file_hash}_{content_hash}"
    
    async def batch_process_documents(self, file_paths: List[str]) -> List[ProcessedDocument]:
        """Process multiple documents concurrently with limits"""
        # Limit concurrent processing
        semaphore = asyncio.Semaphore(3)
        
        async def process_single(file_path: str):
            async with semaphore:
                return await self.process_document(file_path)
        
        tasks = [process_single(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out exceptions
        processed_docs = []
        for result in results:
            if isinstance(result, ProcessedDocument):
                processed_docs.append(result)
            else:
                logger.error(f"Document processing failed: {result}")
        
        return processed_docs 